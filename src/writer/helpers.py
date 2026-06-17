import re

from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from constants import FONT_NAME, FONT_SIZE_MAIN
from omml import latex_to_omml_elem

def set_spacing(para_fmt, before=0, after=0, line=None):
    para_fmt.space_before = Pt(before)
    para_fmt.space_after  = Pt(after)
    if line is not None:
        para_fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        para_fmt.line_spacing = line


def fmt_run(run, size=None, bold=False, italic=False, color=None):
    run.font.name   = FONT_NAME
    run.font.bold   = bold
    run.font.italic = italic
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    rpr    = run._r.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), FONT_NAME)


def set_outline_level(para, level: int):
    pPr = para._p.get_or_add_pPr()
    ol  = pPr.find(qn('w:outlineLvl'))
    if ol is None:
        ol = OxmlElement('w:outlineLvl')
        pPr.append(ol)
    ol.set(qn('w:val'), str(level))


def disable_snap_to_grid(para):
    pPr  = para._p.get_or_add_pPr()
    elem = pPr.find(qn('w:snapToGrid'))
    if elem is None:
        elem = OxmlElement('w:snapToGrid')
        pPr.append(elem)
    elem.set(qn('w:val'), '0')


def set_para_default_font(para, size):
    """Set paragraph default run font size so OMML math objects inherit it."""
    val  = str(int(size.pt * 2))
    pPr  = para._p.get_or_add_pPr()
    rPr  = pPr.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        pPr.append(rPr)
    for tag in ('w:sz', 'w:szCs'):
        elem = rPr.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            rPr.append(elem)
        elem.set(qn('w:val'), val)


def set_keep_with_next(para):
    pPr = para._p.get_or_add_pPr()
    if pPr.find(qn('w:keepNext')) is None:
        pPr.append(OxmlElement('w:keepNext'))


def add_page_number_field(run):
    fld = OxmlElement('w:fldChar')
    fld.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld)
    instr = OxmlElement('w:instrText')
    instr.text = ' PAGE '
    run._r.append(instr)
    fld2 = OxmlElement('w:fldChar')
    fld2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld2)


_INLINE_RE = re.compile(
    r'\*\*(.+?)\*\*'
    r'|\*(.+?)\*'
    r'|`(.+?)`'
    r'|\[(\d+(?:,\s*с\.\s*\d+)?)\]'
    r'|(\$[^$]+?\$)',
    re.DOTALL,
)


def add_inline(para, text: str, base_size=None,
               base_bold: bool = False, base_italic: bool = False):
    if base_size is None:
        base_size = FONT_SIZE_MAIN
    text = text.replace('—', '–')
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            run = para.add_run(text[pos:m.start()])
            fmt_run(run, size=base_size, bold=base_bold, italic=base_italic)
        bold_t, italic_t, code_t, cite_t, math_t = m.groups()
        if bold_t is not None:
            run = para.add_run(bold_t)
            fmt_run(run, size=base_size, bold=True, italic=base_italic)
        elif italic_t is not None:
            run = para.add_run(italic_t)
            fmt_run(run, size=base_size, bold=base_bold, italic=True)
        elif code_t is not None:
            run = para.add_run(code_t)
            run.font.name = 'Courier New'
            run.font.size = base_size
        elif cite_t is not None:
            run = para.add_run(f'[{cite_t}]')
            fmt_run(run, size=base_size, bold=base_bold, italic=base_italic)
        elif math_t is not None:
            omml = latex_to_omml_elem(math_t[1:-1])
            if omml is not None:
                para._p.append(omml)
            else:
                run = para.add_run(math_t[1:-1])
                fmt_run(run, size=base_size, italic=True)
        pos = m.end()
    if pos < len(text):
        run = para.add_run(text[pos:])
        fmt_run(run, size=base_size, bold=base_bold, italic=base_italic)
