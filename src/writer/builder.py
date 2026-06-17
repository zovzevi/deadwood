import base64
import os
import re
import subprocess
import tempfile
import urllib.request
import zlib
from typing import List, Optional

from lxml import etree as _etree

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

try:
    from PIL import Image as PILImage
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

from constants import (
    FONT_SIZE_MAIN, FONT_SIZE_TABLE,
    MARGIN_LEFT, MARGIN_RIGHT, MARGIN_TOP, MARGIN_BOTTOM,
    PARA_INDENT, LINE_SPACING,
)
from omml import latex_to_omml_elem
from writer.helpers import (
    fmt_run, set_spacing, set_outline_level,
    set_keep_with_next, add_page_number_field, add_inline,
    disable_snap_to_grid as _disable_snap_to_grid,
    set_para_default_font as _set_para_default_font,
)


def _prepare_mermaid_code(code: str) -> str:
    # Replace literal \n sequences (backslash + n written in mermaid labels)
    # with <br/> and enable HTML labels so renderers treat them as line breaks.
    code = code.replace('\\n', '<br/>')
    if not code.startswith('%%{'):
        code = "%%{init: {'htmlLabels': true}}%%\n" + code
    return code


def _render_mermaid(code: str) -> Optional[str]:
    """Render Mermaid diagram code to a temporary PNG file.

    Tries mmdc (Mermaid CLI) first, then kroki.io HTTP API.
    Returns the path to the generated PNG, or None on failure.
    The caller is responsible for deleting the file.
    """
    code = _prepare_mermaid_code(code)

    # --- mmdc (Mermaid CLI via Node.js) ---
    try:
        mmd_fd, mmd_path = tempfile.mkstemp(suffix='.mmd')
        png_path = mmd_path[:-4] + '.png'
        try:
            with os.fdopen(mmd_fd, 'w', encoding='utf-8') as f:
                f.write(code)
            result = subprocess.run(
                ['mmdc', '-i', mmd_path, '-o', png_path, '-b', 'white', '--quiet'],
                capture_output=True,
                timeout=30,
            )
            if result.returncode == 0 and os.path.exists(png_path):
                return png_path
        finally:
            if os.path.exists(mmd_path):
                os.unlink(mmd_path)
    except (FileNotFoundError, subprocess.TimeoutExpired, OSError):
        pass

    # --- kroki.io HTTP API ---
    try:
        compressed = zlib.compress(code.encode('utf-8'), 9)
        encoded = base64.urlsafe_b64encode(compressed).decode('ascii')
        url = f'https://kroki.io/mermaid/png/{encoded}'
        req = urllib.request.Request(url, headers={'User-Agent': 'md2docx-gost/1.0'})
        with urllib.request.urlopen(req, timeout=15) as resp:
            data = resp.read()
        png_fd, png_path = tempfile.mkstemp(suffix='.png')
        with os.fdopen(png_fd, 'wb') as f:
            f.write(data)
        return png_path
    except Exception:
        pass

    return None


class GostDocxBuilder:
    def __init__(self, base_dir: str = ''):
        self.base_dir = base_dir
        self.doc = Document()
        body = self.doc.element.body
        for child in list(body.iterchildren(qn('w:p'))):
            body.remove(child)

        self.figure_counter             = 0
        self.table_counter              = 0
        self.formula_counter            = 0
        self.current_appendix           = None
        self.appendix_figure_counter    = 0
        self.appendix_table_counter     = 0
        self.appendix_formula_counter   = 0
        self._last_was_page_break       = False
        self._last_added_blank  = False

        self._setup_page()
        self._setup_math_font()
        self._add_page_numbers()

    def _setup_page(self):
        s               = self.doc.sections[0]
        s.page_width    = Cm(21.0)
        s.page_height   = Cm(29.7)
        s.left_margin   = MARGIN_LEFT
        s.right_margin  = MARGIN_RIGHT
        s.top_margin    = MARGIN_TOP
        s.bottom_margin = MARGIN_BOTTOM

        # Set document grid to "lines" mode so Word stretches word spaces
        # (not inter-character spacing) when justifying text.
        # "linesAndChars" (Word default on some locales) causes letter-spacing expansion.
        sect_pr = s._sectPr
        doc_grid = sect_pr.find(qn('w:docGrid'))
        if doc_grid is None:
            doc_grid = OxmlElement('w:docGrid')
            sect_pr.append(doc_grid)
        doc_grid.set(qn('w:type'), 'lines')
        doc_grid.set(qn('w:linePitch'), '360')

    def _setup_math_font(self):
        _M = "http://schemas.openxmlformats.org/officeDocument/2006/math"
        settings_el = self.doc.settings.element
        math_pr = settings_el.find(f'{{{_M}}}mathPr')
        if math_pr is None:
            math_pr = _etree.SubElement(settings_el, f'{{{_M}}}mathPr')
        for child_tag in (f'{{{_M}}}mathFont',):
            child = math_pr.find(child_tag)
            if child is None:
                child = _etree.SubElement(math_pr, child_tag)
            child.set(f'{{{_M}}}val', 'Times New Roman')

    def _add_page_numbers(self):
        footer = self.doc.sections[0].footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.clear()
        fp            = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment  = WD_ALIGN_PARAGRAPH.CENTER
        run           = fp.add_run()
        add_page_number_field(run)
        fmt_run(run, size=FONT_SIZE_MAIN)

    def _base_para_fmt(self, para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                       first_line=None, left_indent=None):
        fmt           = para.paragraph_format
        fmt.alignment = align
        set_spacing(fmt, before=0, after=0, line=LINE_SPACING)
        if first_line  is not None:
            fmt.first_line_indent = first_line
        if left_indent is not None:
            fmt.left_indent = left_indent
        _disable_snap_to_grid(para)
        _set_para_default_font(para, FONT_SIZE_MAIN)
        return fmt

    def add_structural_heading(self, text: str, page_break_before: bool = True):
        if page_break_before:
            self._page_break()
        self._last_was_page_break      = False
        self._last_added_blank = False
        para                          = self.doc.add_paragraph()
        fmt                           = self._base_para_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER)
        fmt.first_line_indent         = Pt(0)
        fmt.space_before = fmt.space_after = Pt(0)
        run = para.add_run(text.upper())
        fmt_run(run, size=FONT_SIZE_MAIN, bold=True)
        if text.upper() != 'СОДЕРЖАНИЕ':
            set_outline_level(para, 0)
        set_keep_with_next(para)
        self._blank_line()
        self._last_added_blank = True
        return para

    def add_section_heading(self, number: str, text: str, level: int = 1):
        if level == 1 and number:
            self._page_break()
            self._last_added_blank = False
        elif level >= 2:
            if not self._last_added_blank:
                self._blank_line()
        self._last_was_page_break      = False
        self._last_added_blank = False
        full_text                     = f'{number} {text}'.strip() if number else text
        para                          = self.doc.add_paragraph()
        fmt                           = self._base_para_fmt(para, align=WD_ALIGN_PARAGRAPH.LEFT)
        fmt.first_line_indent         = PARA_INDENT
        fmt.space_before = fmt.space_after = Pt(0)
        run = para.add_run(full_text)
        fmt_run(run, size=FONT_SIZE_MAIN, bold=(level <= 2))
        if level <= 3:
            set_outline_level(para, level - 1)
        set_keep_with_next(para)
        self._blank_line()
        self._last_added_blank = True
        return para

    def add_appendix_heading(self, letter: str):
        self.current_appendix         = letter
        self.appendix_figure_counter  = 0
        self.appendix_table_counter   = 0
        self.appendix_formula_counter = 0
        self._page_break()
        self._last_was_page_break = False
        para                          = self.doc.add_paragraph()
        fmt                           = self._base_para_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER)
        fmt.first_line_indent         = Pt(0)
        run = para.add_run(f'ПРИЛОЖЕНИЕ {letter}')
        fmt_run(run, size=FONT_SIZE_MAIN, bold=True)
        set_outline_level(para, 0)
        set_keep_with_next(para)
        return para

    def add_appendix_title(self, text: str):
        self._last_was_page_break = False
        para                  = self.doc.add_paragraph()
        fmt                   = self._base_para_fmt(para, align=WD_ALIGN_PARAGRAPH.CENTER)
        fmt.first_line_indent = Pt(0)
        add_inline(para, text, base_size=FONT_SIZE_MAIN)
        set_keep_with_next(para)
        self._blank_line()
        return para

    def add_paragraph(self, text: str):
        self._last_was_page_break      = False
        self._last_added_blank = False
        para = self.doc.add_paragraph()
        self._base_para_fmt(para, first_line=PARA_INDENT)
        add_inline(para, text)
        return para

    def add_list_item(self, marker_type: str, marker: str,
                      text: str, indent_level: int = 0):
        self._last_was_page_break      = False
        self._last_added_blank = False
        para                  = self.doc.add_paragraph()
        fmt                   = para.paragraph_format
        fmt.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
        fmt.space_after       = Pt(0)
        base                  = PARA_INDENT + Cm(indent_level * 1.25)
        fmt.left_indent       = base + Cm(0.5)
        fmt.first_line_indent = Cm(-0.5)
        set_spacing(fmt, before=0, after=0, line=LINE_SPACING)
        _disable_snap_to_grid(para)
        fmt_run(para.add_run(marker + '\t'), size=FONT_SIZE_MAIN)
        add_inline(para, text)
        return para

    def add_table(self, header, rows, caption=None):
        self._last_was_page_break = False
        if not self._last_added_blank:
            self._blank_line()
        self._last_added_blank = False
        cap_text = self._resolve_table_caption(caption)

        cap_para                      = self.doc.add_paragraph()
        fmt                           = self._base_para_fmt(cap_para, align=WD_ALIGN_PARAGRAPH.LEFT)
        fmt.first_line_indent         = Pt(0)
        fmt.space_after               = Pt(3)
        add_inline(cap_para, cap_text, base_size=FONT_SIZE_MAIN)
        set_keep_with_next(cap_para)

        col_count = max(
            len(header) if header else 0,
            max((len(r) for r in rows), default=0),
        )
        if col_count == 0:
            return None

        tbl           = self.doc.add_table(rows=0, cols=col_count)
        tbl.style     = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        if header:
            hr = tbl.add_row()
            for j, ct in enumerate(header[:col_count]):
                cell = hr.cells[j]
                cell.text = ''
                p         = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline(p, ct, base_size=FONT_SIZE_TABLE, base_bold=True)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for row_data in rows:
            r = tbl.add_row()
            for j, ct in enumerate(row_data[:col_count]):
                cell        = r.cells[j]
                cell.text   = ''
                p           = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                txt = (ct or '').strip() or '–'
                if txt == '-':
                    txt = '–'
                add_inline(p, txt, base_size=FONT_SIZE_TABLE)

        self._blank_line()
        self._last_added_blank = True
        return tbl

    def _resolve_table_caption(self, raw: Optional[str]) -> str:
        if raw and re.match(r'^Таблица\s+[\w.]+\s*–', raw):
            return raw
        if self.current_appendix:
            self.appendix_table_counter += 1
            num = f'{self.current_appendix}.{self.appendix_table_counter}'
        else:
            self.table_counter += 1
            num = str(self.table_counter)
        return f'Таблица {num}' + (f' – {raw}' if raw else '')

    def add_code_block(self, code: str, language: str = ''):
        self._last_was_page_break = False
        if not self._last_added_blank:
            self._blank_line()
        self._last_added_blank = False

        lines = code.split('\n')
        for line in lines:
            para = self.doc.add_paragraph()
            fmt = para.paragraph_format
            fmt.alignment = WD_ALIGN_PARAGRAPH.LEFT
            fmt.first_line_indent = Pt(0)
            fmt.left_indent = PARA_INDENT
            set_spacing(fmt, before=0, after=0, line=1.0)
            _disable_snap_to_grid(para)

            pPr = para._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'F2F2F2')
            pPr.append(shd)

            run = para.add_run(line if line else ' ')
            run.font.name = 'Courier New'
            run.font.size = FONT_SIZE_TABLE
            rpr = run._r.get_or_add_rPr()
            rFonts = rpr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rpr.insert(0, rFonts)
            for attr in ('w:ascii', 'w:hAnsi', 'w:cs'):
                rFonts.set(qn(attr), 'Courier New')

        self._blank_line()
        self._last_added_blank = True

    def add_figure(self, src: str, alt: str):
        self._last_was_page_break = False
        if not self._last_added_blank:
            self._blank_line()
        self._last_added_blank = False
        if self.current_appendix:
            self.appendix_figure_counter += 1
            num = f'{self.current_appendix}.{self.appendix_figure_counter}'
        else:
            self.figure_counter += 1
            num = str(self.figure_counter)

        resolved = os.path.join(self.base_dir, src) if self.base_dir else src
        img_added = False
        if src and os.path.exists(resolved):
            try:
                para           = self.doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.add_run().add_picture(resolved, width=Cm(15.5))
                img_added = True
            except Exception:
                pass

        if not img_added:
            para           = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt_run(para.add_run(f'[Рисунок: {src}]'), size=FONT_SIZE_MAIN, italic=True)

        cap_para                      = self.doc.add_paragraph()
        fmt                           = self._base_para_fmt(cap_para, align=WD_ALIGN_PARAGRAPH.CENTER)
        fmt.first_line_indent         = Pt(0)
        caption_text = f'Рисунок {num}' + (f' – {alt}' if alt else '')
        add_inline(cap_para, caption_text, base_size=FONT_SIZE_MAIN)
        self._blank_line()
        self._last_added_blank = True
        return cap_para

    def add_diagram(self, code: str, alt: str):
        self._last_was_page_break = False
        if not self._last_added_blank:
            self._blank_line()
        self._last_added_blank = False
        if self.current_appendix:
            self.appendix_figure_counter += 1
            num = f'{self.current_appendix}.{self.appendix_figure_counter}'
        else:
            self.figure_counter += 1
            num = str(self.figure_counter)

        tmp_png = _render_mermaid(code)
        img_added = False

        if tmp_png:
            try:
                para           = self.doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.add_run().add_picture(tmp_png, width=Cm(15.5))
                img_added = True
            except Exception:
                pass
            finally:
                try:
                    os.unlink(tmp_png)
                except OSError:
                    pass

        if not img_added:
            para           = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt_run(para.add_run(f'[Диаграмма: {alt}]'), size=FONT_SIZE_MAIN, italic=True)

        cap_para              = self.doc.add_paragraph()
        fmt                   = self._base_para_fmt(cap_para, align=WD_ALIGN_PARAGRAPH.CENTER)
        fmt.first_line_indent = Pt(0)
        caption_text = f'Рисунок {num}' + (f' – {alt}' if alt else '')
        add_inline(cap_para, caption_text, base_size=FONT_SIZE_MAIN)
        self._blank_line()
        self._last_added_blank = True
        return cap_para

    def add_formula(self, formula_text: str, explicit_number: Optional[str] = None):
        self._last_was_page_break = False
        if self.current_appendix:
            self.appendix_formula_counter += 1
            num = f'{self.current_appendix}.{self.appendix_formula_counter}'
        else:
            self.formula_counter += 1
            num = explicit_number or str(self.formula_counter)

        if not self._last_added_blank:
            self._blank_line()
        self._last_added_blank = False
        para                  = self.doc.add_paragraph()
        fmt                   = para.paragraph_format
        fmt.alignment         = WD_ALIGN_PARAGRAPH.CENTER
        fmt.first_line_indent = Pt(0)
        set_spacing(fmt, before=0, after=0, line=LINE_SPACING)
        _disable_snap_to_grid(para)
        _set_para_default_font(para, FONT_SIZE_MAIN)

        omml = latex_to_omml_elem(formula_text)
        if omml is not None:
            para._p.append(omml)
        else:
            fmt_run(para.add_run(formula_text), size=FONT_SIZE_MAIN, italic=True)

        fmt_run(para.add_run('\t'),     size=FONT_SIZE_MAIN)
        fmt_run(para.add_run(f'({num})'), size=FONT_SIZE_MAIN)

        pPr       = para._p.get_or_add_pPr()
        tabs_elem = pPr.find(qn('w:tabs'))
        if tabs_elem is None:
            tabs_elem = OxmlElement('w:tabs')
            pPr.append(tabs_elem)
        tab_stop = OxmlElement('w:tab')
        tab_stop.set(qn('w:val'), 'right')
        tab_stop.set(qn('w:pos'), str(int(15.5 * 567)))  # 15.5 cm → twips
        tabs_elem.append(tab_stop)

        self._blank_line()
        self._last_added_blank = True
        return para

    def add_formula_where(self, items: List[str]):
        for idx, item in enumerate(items):
            para                  = self.doc.add_paragraph()
            fmt                   = para.paragraph_format
            fmt.alignment         = WD_ALIGN_PARAGRAPH.JUSTIFY
            fmt.first_line_indent = Pt(0)
            fmt.left_indent       = PARA_INDENT
            set_spacing(fmt, before=0, after=0, line=LINE_SPACING)
            clean = item.strip().rstrip('.,;')
            if idx == 0:
                fmt_run(para.add_run('где '), size=FONT_SIZE_MAIN)
            add_inline(para, clean, base_size=FONT_SIZE_MAIN)
            fmt_run(para.add_run(';' if idx < len(items) - 1 else '.'), size=FONT_SIZE_MAIN)
        self._blank_line()
        self._last_added_blank = True

    def add_toc(self):
        para           = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run            = para.add_run()

        for fld_type in ('begin', 'separate', 'end'):
            if fld_type == 'begin':
                fld = OxmlElement('w:fldChar')
                fld.set(qn('w:fldCharType'), 'begin')
                run._r.append(fld)
                instr = OxmlElement('w:instrText')
                instr.set(qn('xml:space'), 'preserve')
                instr.text = r' TOC \o "1-3" \h \z \u '
                run._r.append(instr)
            else:
                fld = OxmlElement('w:fldChar')
                fld.set(qn('w:fldCharType'), fld_type)
                run._r.append(fld)

        fmt_run(run, size=FONT_SIZE_MAIN)
        self._blank_line()

    def _page_break(self):
        if self._last_was_page_break:
            return
        self._last_was_page_break = True
        para = self.doc.add_paragraph()
        para.add_run().add_break(WD_BREAK.PAGE)
        return para

    def _blank_line(self):
        para      = self.doc.add_paragraph()
        run       = para.add_run('')
        run.font.size = FONT_SIZE_MAIN
        set_spacing(para.paragraph_format, before=0, after=0, line=LINE_SPACING)
        return para

    def save(self, path: str):
        self.doc.save(path)
